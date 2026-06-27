"""Generate the Greek technical report (.docx) for Wine DSS, max 4 pages, no cover."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
DIAG = ROOT / "diagrams2"
OUT = ROOT / "docs" / "wine-dss-final-report.docx"

doc = Document()

# --- compact page + base style ---
sec = doc.sections[0]
sec.top_margin = Cm(1.4); sec.bottom_margin = Cm(1.4)
sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.04

ACCENT = RGBColor(0x6B, 0x10, 0x2A)  # wine red

def set_heading(level, size, space_before):
    st = doc.styles[f"Heading {level}"]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = ACCENT
    st.paragraph_format.space_before = Pt(space_before)
    st.paragraph_format.space_after = Pt(2)

set_heading(1, 13, 6)
set_heading(2, 11, 4)

def h(text, level=1):
    doc.add_heading(text, level=level)

def p(text, bold_lead=None):
    par = doc.add_paragraph()
    if bold_lead:
        r = par.add_run(bold_lead); r.bold = True
    par.add_run(text)
    return par

def bullet(text, lead=None):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(1)
    par.paragraph_format.left_indent = Cm(0.6)
    if lead:
        r = par.add_run(lead); r.bold = True
    par.add_run(text)
    return par

def img(name, width_cm, caption):
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(2); par.paragraph_format.space_after = Pt(1)
    par.add_run().add_picture(str(DIAG / name), width=Cm(width_cm))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x55,0x55,0x55)

# ===================== TITLE =====================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Wine DSS — Σύστημα Υποστήριξης Αποφάσεων για Ζευγάρωμα Κρασιού & Φαγητού")
r.bold = True; r.font.size = Pt(15); r.font.color.rgb = ACCENT
t.paragraph_format.space_after = Pt(1)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Τεχνική Αναφορά — Οντολογία (OWL/HermiT) + Generative AI μέσω MCP")
rs.italic = True; rs.font.size = Pt(10); rs.font.color.rgb = RGBColor(0x55,0x55,0x55)
sub.paragraph_format.space_after = Pt(6)

# ===================== 1. ΕΙΣΑΓΩΓΗ =====================
h("1. Εισαγωγή & Πρόβλημα")
p("Το Wine DSS είναι ένα σημασιολογικό Σύστημα Υποστήριξης Αποφάσεων που αυτοματοποιεί "
  "τη σύσταση κρασιών για ένα εστιατόριο/κάβα. Το πρόβλημα που λύνει: δοθέντος ενός πιάτου "
  "ή μιας κατηγορίας φαγητού, ποιες φιάλες από το διαθέσιμο stock ταιριάζουν καλύτερα — και "
  "γιατί. Η καινοτομία είναι ότι οι κανόνες ζευγαρώματος δεν είναι «καρφωμένοι» στον κώδικα, "
  "αλλά εκφράζονται δηλωτικά ως αξιώματα οντολογίας OWL και εξάγονται αυτόματα από έναν reasoner. "
  "Το σύστημα εκτίθεται ως MCP server, ώστε ένα LLM να το χρησιμοποιεί συνομιλιακά (μοτίβο RAG / "
  "Generative AI — Ιδέα 4 της εκφώνησης).")

# ===================== 2. ΑΡΧΙΤΕΚΤΟΝΙΚΗ =====================
h("2. Αρχιτεκτονική & Ανάπτυξη")
p("Η αρχιτεκτονική είναι ένας αγωγός (pipeline) τριών επιπέδων. «Εγκέφαλος» που ενορχηστρώνει "
  "τη ροή είναι το Python script (Recommender). Η οντολογία δεν εκτελεί κώδικα· λειτουργεί "
  "αποκλειστικά ως η έξυπνη Βάση Γνώσης.")
bullet("ορίζει 9 εργαλεία (MCP tools) που καλεί το LLM-πελάτης μέσω streamable-HTTP (θύρα 8000).",
       lead="Επίπεδο Παρουσίασης — FastMCP Server (server.py): ")
bullet("ο Recommender συντονίζει Βάση Γνώσης + Κατάλογο και επιστρέφει ενοποιημένο αποτέλεσμα.",
       lead="Επίπεδο Λογικής — Ενορχηστρωτής (recommender.py): ")
bullet("η KnowledgeBase (owlready2 + HermiT, ερωτήματα SPARQL) για τη γνώση ζευγαρώματος, "
       "και ο Catalogue (SQLite) για τα 221 κρασιά & 21 πιάτα του καταστήματος.",
       lead="Επίπεδο Δεδομένων: ")

p("Η διανομή γίνεται ως ένα αυτοτελές Docker image. Το Dockerfile ξεκινά από δύο βασικές εικόνες "
  "του Docker Hub — python:3.12-slim και eclipse-temurin:17-jre (το JRE απαιτείται από τον HermiT) "
  "— και παράγει μία εικόνα wine-dss που περιέχει τις εξαρτήσεις, τον κώδικα (src) και τα δεδομένα "
  "(data). Με «docker compose up» το daemon χτίζει την εικόνα και τρέχει έναν container που εκθέτει "
  "τη θύρα 8000· δεν απαιτείται καμία εγκατάσταση στον host.")
img("01_docker.png", 7.5, "Σχήμα 1. Αρχιτεκτονική Docker: το docker compose χτίζει την εικόνα wine-dss και τρέχει τον container στη θύρα 8000.")
img("02_networking.png", 16.6, "Σχήμα 2. Δικτύωση: ο MCP client συνδέεται μέσω HTTP στη θύρα 8000 του host, που προωθείται στον container (0.0.0.0:8000/mcp).")
img("03_server_internals.png", 7.5, "Σχήμα 3. Εσωτερική αρχιτεκτονική του server: FastMCP, 9 εργαλεία, Recommender, KnowledgeBase και Catalogue.")

# ===================== 3. ΜΟΝΤΕΛΟΠΟΙΗΣΗ ΓΝΩΣΗΣ =====================
h("3. Μοντελοποίηση Γνώσης (Οντολογία)")
p("Η οντολογία (wine_dss.owl) περιγράφει ποικιλίες σταφυλιού και τα οργανοληπτικά τους "
  "χαρακτηριστικά ως ελεγχόμενο λεξιλόγιο (named individuals), όχι ως ελεύθερους αριθμούς.")
bullet("WineVariety και υποκλάσεις ανά χρώμα/σώμα/ζάχαρη (RedVariety, WhiteVariety, "
       "DryVariety) και 7 κλάσεις ζευγαρώματος (SteakPairing, SeafoodPairing, "
       "CheesePairing, DessertPairing, SpicyFoodPairing, RosePairing, PoultryPairing).", lead="Κλάσεις: ")
bullet("Object properties hasTypicalColor, canProduceColor, hasTypicalBody, hasTypicalSugar, "
       "hasTypicalFlavor, originatesFrom· data properties hasVarietyName, hasSKU, hasSynonym, "
       "isIndigenousToGreece, isAromatic.", lead="Ιδιότητες: ")
bullet("75 ποικιλίες σταφυλιού (π.χ. Xinomavro V-064, Agiorgitiko), συνδεδεμένες με τιμές "
       "χρώματος/σώματος/ζάχαρης και χώρες προέλευσης.", lead="Στιγμιότυπα: ")
bullet("κάθε κλάση ζευγαρώματος ορίζεται με owl:equivalentClass ως τομή περιορισμών (hasValue). "
       "Ο HermiT reasoner ταξινομεί αυτόματα κάθε ποικιλία στις κλάσεις που ικανοποιεί — αυτή "
       "είναι η εξαγωγή έμμεσης γνώσης (inference).", lead="Κανόνες/Αξιώματα: ")

# ===================== 4. ΛΟΓΙΚΗ ΤΑΙΡΙΑΣΜΑΤΟΣ =====================
h("4. Η Λογική του Ταιριάσματος (Pairing)")
p("Το ζευγάρωμα προκύπτει σε δύο φάσεις: (α) σημασιολογική εξαγωγή των κατάλληλων ποικιλιών "
  "από τον reasoner, και (β) φιλτράρισμα του πραγματικού καταλόγου. Το κρίσιμο σημείο είναι ότι "
  "ο κανόνας ζευγαρώματος ζει στην οντολογία ως ορισμός κλάσης:")
# rule table-like list
rules = [
    ("SteakPairing", "Color = Red  ⊓  Body = Full"),
    ("SeafoodPairing", "Color = White  ⊓  Body = Light"),
    ("CheesePairing", "Body = Full"),
    ("DessertPairing", "Sugar = Sweet"),
    ("SpicyFoodPairing", "Color = White  ⊓  Sugar = OffDry"),
    ("RosePairing", "canProduceColor = Rose"),
    ("PoultryPairing", "Color = Red  ⊓  Body = Light"),
]
for cls, rule in rules:
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(0); par.paragraph_format.left_indent = Cm(0.6)
    rr = par.add_run(f"{cls} ≡ "); rr.bold = True; rr.font.size = Pt(9.5)
    rc = par.add_run(rule); rc.font.name = "Consolas"; rc.font.size = Pt(9.5)

img("04_dataflow.png", 16.6, "Σχήμα 4. Ροή δεδομένων: από το dish ή food_type, η KnowledgeBase δίνει ποικιλίες, ο Catalogue τις φιάλες, και ακολουθεί βαθμολόγηση και κατάταξη.")

p("Επειδή το Xinomavro έχει Color=Red, Body=Full και canProduceColor=Rose, ο reasoner το "
  "κατατάσσει ταυτόχρονα σε SteakPairing, CheesePairing και RosePairing — χωρίς καμία γραμμή "
  "if/else. Η ροή εκτέλεσης για ένα ερώτημα είναι:")
bullet("Ο χρήστης ζητά κρασί για ένα πιάτο· το Catalogue.lookup_dish επιστρέφει το food_type "
       "(π.χ. «steak» → SteakPairing).", lead="1. Ανάκτηση: ")
bullet("Η KnowledgeBase τρέχει τον reasoner (μία φορά, με caching) και με SPARQL πάνω στο "
       "reasoned graph επιστρέφει τις ποικιλίες της κλάσης (και των υποκλάσεών της).",
       lead="2. Εξαγωγή: ")
bullet("το Catalogue φιλτράρει το stock (variety IN (...), max_price, color, σε απόθεμα) και "
       "ταξινομεί κατά τιμή φιάλης.", lead="3. Φιλτράρισμα: ")
bullet("το LLM-πελάτης λαμβάνει δομημένο αποτέλεσμα {dish, food_type, varieties, wines} και "
       "συνθέτει αιτιολογημένη σύσταση.", lead="4. Παρουσίαση: ")
img("05_sequence.png", 16.6, "Σχήμα 5. Ακολουθία κλήσεων μιας σύστασης, από το αίτημα του χρήστη μέχρι τις τελικές φιάλες.")

# ===================== 5. ΕΠΕΚΤΑΣΗ =====================
h("5. Επέκταση: Generative AI / RAG μέσω MCP")
p("Η δεύτερη μεθοδολογία (Μέρος Β) είναι η Παραγωγική ΤΝ με μοτίβο RAG. Αντί για ad-hoc prompts, "
  "η γνώση εκτίθεται ως τυποποιημένα εργαλεία Model Context Protocol (list_dishes, search_dishes, "
  "list_categories, varieties_in_category, recommend_wine_for_dish, recommend_wine_for_food_type, "
  "classify_variety, variety_info, search_catalogue). Το LLM ανακαλεί τεκμηριωμένη γνώση από την οντολογία (grounding) και την "
  "μετατρέπει σε φυσική, εξηγήσιμη σύσταση προς τον sommelier — π.χ. «Για το steak προτείνω Xinomavro "
  "γιατί ταξινομείται ως SteakPairing (κόκκινο, γεμάτο σώμα)· διαθέσιμη φιάλη εντός προϋπολογισμού…».")

# ===================== 6. ΜΕΛΛΟΝΤΙΚΕΣ ΒΕΛΤΙΩΣΕΙΣ =====================
h("6. Μελλοντικές Βελτιώσεις")
bullet("ενσωμάτωση Fuzzy Inference System (scikit-fuzzy) ώστε αριθμητικά χαρακτηριστικά "
       "(τιμή, αξιολόγηση) να παράγουν ασαφές «Value-for-Money / Pairing Score» αντί για δυαδική ένταξη.",
       lead="Ασαφής βαθμολόγηση: ")
bullet("πρόβλεψη ζήτησης/αναπλήρωσης stock με Random Forest πάνω σε ιστορικά δεδομένα πωλήσεων, "
       "συνδυασμένη με την κατηγοριοποίηση της οντολογίας (XAI με SHAP για εξήγηση).",
       lead="Machine Learning: ")
bullet("πρόβλημα Knapsack/βελτιστοποίηση καρτών κρασιού με Γενετικό Αλγόριθμο (PyGAD) υπό περιορισμό "
       "προϋπολογισμού/ποικιλίας.", lead="Βελτιστοποίηση: ")
bullet("έλεγχος διαθεσιμότητας Java στην εκκίνηση, ασύγχρονη εκτέλεση HermiT με timeout, structured "
       "logging/health-checks, test suite και επικύρωση σχήματος DB.", lead="Στιβαρότητα μηχανικής: ")
bullet("πλουσιότεροι κανόνες SWRL (π.χ. ταξινόμηση «PremiumPairing» βάσει συνδυασμού περιοχής & σώματος) "
       "και επέκταση πέρα από τα ελληνικά σταφύλια.", lead="Εμπλουτισμός γνώσης: ")

# ===================== 7. ΣΥΜΠΕΡΑΣΜΑΤΑ =====================
h("7. Συμπεράσματα")
p("Το Wine DSS δείχνει πώς μια οντολογία OWL με reasoner μπορεί να αποτελέσει τον δηλωτικό πυρήνα "
  "αποφάσεων, ενώ ένα LLM μέσω MCP προσφέρει το συνομιλιακό, εξηγήσιμο περίβλημα. Ο διαχωρισμός "
  "κανόνων (οντολογία) από ενορχήστρωση (Python) καθιστά το σύστημα επεκτάσιμο και συντηρήσιμο: "
  "νέοι κανόνες ζευγαρώματος προστίθενται ως αξιώματα, χωρίς αλλαγή κώδικα.")

OUT.parent.mkdir(exist_ok=True)
doc.save(str(OUT))
print("saved:", OUT)
